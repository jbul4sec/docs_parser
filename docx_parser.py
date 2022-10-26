from docx import Document


class DocxParser:
    """Class for docx parsing."""

    @classmethod
    def get_title_or_description(cls, path_to_file: str) -> tuple:
        """
        Get title and description from Word file
        :param path_to_file: path to file
        :return: tuple with title and description from a file.
        """
        title = ''
        description = ''
        doc = Document(path_to_file)

        for paragraph in doc.paragraphs:
            if paragraph.text != '':
                title += paragraph.text
                break

        for paragraph in doc.paragraphs[1:40]:
            description += paragraph.text + ' '

        return title, description


class DocxTable:
    """Create a table and save it in word file."""

    def __init__(self, columns: int, filename: str):
        self.rows = 1
        self.columns = columns
        self.filename = filename
        self.doc_table = Document()
        self.table_handler = self.create_table()

    def create_table(self):
        return self.doc_table.add_table(rows=self.rows, cols=self.columns)

    def get_table(self):
        return self.table_handler

    def save_table(self):
        self.doc_table.save(f'{self.filename}.docx')


def fulfil_table(number: int, file: str, title: str, description: str, table: Document):
    if number == 1:
        row_cells = table.rows[0].cells
        row_cells[0].text = '№'
        row_cells[1].text = 'Имя файла'
        row_cells[2].text = 'Титул'
        row_cells[3].text = 'Краткое описание'
    else:
        row_cells = table.add_row().cells
        row_cells[0].text = str(number)
        row_cells[1].text = file
        row_cells[2].text = title
        row_cells[3].text = description
